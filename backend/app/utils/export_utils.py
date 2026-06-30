"""
Export utilities for summary output.
Generates PDF, DOCX, Markdown, and plain text formats.
"""
import io
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def export_as_txt(data: dict) -> bytes:
    lines = []
    if data.get("title"):
        lines += [data["title"], "=" * len(data["title"]), ""]
    lines += [f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]
    if data.get("summary"):
        lines += ["SUMMARY", "-------", data["summary"], ""]
    if data.get("bullets"):
        lines += ["KEY POINTS", "----------"]
        lines += [f"• {b}" for b in data["bullets"]]
        lines.append("")
    if data.get("takeaways"):
        lines += ["KEY TAKEAWAYS", "-------------"]
        lines += [f"• {t}" for t in data["takeaways"]]
        lines.append("")
    if data.get("actions"):
        lines += ["ACTION ITEMS", "------------"]
        lines += [f"☐ {a}" for a in data["actions"]]
        lines.append("")
    meta = []
    if data.get("wordCount"):
        meta.append(f"Word count: {data['wordCount']}")
    if data.get("readingTime"):
        meta.append(f"Reading time: {data['readingTime']} min")
    if data.get("compression"):
        meta.append(f"Compression: {data['compression']}")
    if data.get("sentiment"):
        meta.append(f"Sentiment: {data['sentiment']}")
    if meta:
        lines += ["METADATA", "--------"] + meta
    return "\n".join(lines).encode("utf-8")


def export_as_markdown(data: dict) -> bytes:
    lines = []
    if data.get("title"):
        lines += [f"# {data['title']}", ""]
    lines.append(f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n")
    if data.get("summary"):
        lines += ["## Summary", "", data["summary"], ""]
    if data.get("bullets"):
        lines += ["## Key Points", ""]
        lines += [f"- {b}" for b in data["bullets"]]
        lines.append("")
    if data.get("takeaways"):
        lines += ["## Key Takeaways", ""]
        lines += [f"- {t}" for t in data["takeaways"]]
        lines.append("")
    if data.get("actions"):
        lines += ["## Action Items", ""]
        lines += [f"- [ ] {a}" for a in data["actions"]]
        lines.append("")
    meta_parts = []
    for key, label in [("wordCount", "Words"), ("readingTime", "Reading time"), ("compression", "Compression"), ("sentiment", "Sentiment"), ("difficulty", "Difficulty")]:
        if data.get(key):
            val = f"{data[key]} min" if key == "readingTime" else str(data[key])
            meta_parts.append(f"**{label}:** {val}")
    if meta_parts:
        lines += ["## Metadata", ""] + meta_parts
    return "\n".join(lines).encode("utf-8")


def export_as_pdf(data: dict) -> bytes:
    """Generate a neatly styled PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        buffer = io.BytesIO()
        # Clean 0.75-inch (54 points) margins
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=54, leftMargin=54,
                                topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        
        # Color palette
        primary_color = colors.HexColor('#7c6af7')    # SummrAI Purple
        secondary_color = colors.HexColor('#f472b6')  # Pink Accent
        text_dark = colors.HexColor('#1e1e24')        # Charcoal text
        text_light = colors.HexColor('#4a4a50')       # Slate subtext
        bg_card = colors.HexColor('#f8f7ff')          # Very light purple/gray card
        border_card = colors.HexColor('#e6e4f9')      # Soft border
        
        # Custom styles
        title_style = ParagraphStyle(
            'PdfTitle', parent=styles['Heading1'],
            textColor=primary_color, fontSize=22, spaceAfter=4, leading=26
        )
        subtitle_style = ParagraphStyle(
            'PdfSubtitle', parent=styles['Normal'],
            textColor=text_light, fontSize=10, spaceAfter=14
        )
        h2_style = ParagraphStyle(
            'PdfH2', parent=styles['Heading2'],
            textColor=primary_color, fontSize=14, spaceBefore=18, spaceAfter=8,
            keepWithNext=True
        )
        body_style = ParagraphStyle(
            'PdfBody', parent=styles['Normal'],
            textColor=text_dark, fontSize=10.5, leading=16, spaceAfter=10
        )
        bullet_style = ParagraphStyle(
            'PdfBullet', parent=styles['Normal'],
            textColor=text_dark, fontSize=10, leading=15,
            leftIndent=15, firstLineIndent=-10, spaceAfter=6
        )
        meta_label_style = ParagraphStyle(
            'PdfMetaLabel', parent=styles['Normal'],
            textColor=text_light, fontSize=9, alignment=TA_CENTER
        )
        meta_value_style = ParagraphStyle(
            'PdfMetaVal', parent=styles['Normal'],
            textColor=primary_color, fontSize=10.5, leading=12, alignment=TA_CENTER
        )
        footer_style = ParagraphStyle(
            'PdfFooter', parent=styles['Normal'],
            textColor=colors.HexColor('#a0a0b0'), fontSize=8, alignment=TA_CENTER
        )

        story = []
        
        # Header title
        title_text = data.get("title") or "SummrAI Document Summary"
        story.append(Paragraph(title_text, title_style))
        
        gen_time = datetime.now().strftime('%B %d, %Y at %I:%M %p')
        story.append(Paragraph(f"Summarized on {gen_time} &bull; Generated by SummrAI", subtitle_style))
        story.append(HRFlowable(color=primary_color, thickness=1.5, spaceAfter=14))

        # Metadata Card (Table)
        meta_cols = []
        if data.get("wordCount"):
            meta_cols.append([Paragraph("WORDS", meta_label_style), Paragraph(str(data["wordCount"]), meta_value_style)])
        if data.get("readingTime"):
            meta_cols.append([Paragraph("READ TIME", meta_label_style), Paragraph(f"{data['readingTime']}m", meta_value_style)])
        if data.get("compression"):
            meta_cols.append([Paragraph("COMPRESSION", meta_label_style), Paragraph(str(data["compression"]), meta_value_style)])
        if data.get("sentiment"):
            meta_cols.append([Paragraph("SENTIMENT", meta_label_style), Paragraph(str(data["sentiment"]).capitalize(), meta_value_style)])
        if data.get("difficulty"):
            meta_cols.append([Paragraph("COMPLEXITY", meta_label_style), Paragraph(str(data["difficulty"]).capitalize(), meta_value_style)])

        if meta_cols:
            headers = [col[0] for col in meta_cols]
            values = [col[1] for col in meta_cols]
            
            table_data = [headers, values]
            # Width calculation: 504 pt available (612 page width - 108 margins)
            col_width = 504 / len(meta_cols)
            
            t = Table(table_data, colWidths=[col_width] * len(meta_cols))
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), bg_card),
                ('BOX', (0,0), (-1,-1), 0.5, border_card),
                ('INNERGRID', (0,0), (-1,-1), 0.25, border_card),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ]))
            story.append(t)
            story.append(Spacer(1, 15))

        # 1. Summary Section
        if data.get("summary"):
            story.append(Paragraph("Executive Summary", h2_style))
            story.append(Paragraph(data["summary"], body_style))
            story.append(Spacer(1, 4))

        # 2. Key Points
        if data.get("bullets") and len(data["bullets"]) > 0:
            story.append(Paragraph("Key Points", h2_style))
            for b in data["bullets"]:
                if b.strip():
                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{b}", bullet_style))
            story.append(Spacer(1, 4))

        # 3. Key Takeaways
        if data.get("takeaways") and len(data["takeaways"]) > 0:
            story.append(Paragraph("Key Takeaways", h2_style))
            for t_item in data["takeaways"]:
                if t_item.strip():
                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{t_item}", bullet_style))
            story.append(Spacer(1, 4))

        # 4. Action Items
        if data.get("actions") and len(data["actions"]) > 0:
            story.append(Paragraph("Action Items", h2_style))
            for a in data["actions"]:
                if a.strip():
                    story.append(Paragraph(f"&#9633;&nbsp;&nbsp;{a}", bullet_style))
            story.append(Spacer(1, 4))

        # 5. Sentiment Analysis Note
        if data.get("sentimentNote"):
            story.append(Paragraph("Sentiment Context", h2_style))
            story.append(Paragraph(data["sentimentNote"], body_style))

        # 6. One-Sentence TLDR
        if data.get("tldr"):
            story.append(Paragraph("One-Sentence TLDR", h2_style))
            story.append(Paragraph(f"<i>\"{data['tldr']}\"</i>", body_style))
            story.append(Spacer(1, 4))

        # 7. Questions & Quiz
        if data.get("questions") and len(data["questions"]) > 0:
            story.append(Paragraph("Questions & Quiz", h2_style))
            for q_str in data["questions"]:
                story.append(Paragraph(q_str, bullet_style))
            story.append(Spacer(1, 4))

        # 8. Mind Map
        if data.get("mindmap"):
            story.append(Paragraph("Mind Map Tree", h2_style))
            mindmap_style = ParagraphStyle(
                'PdfMindmap', parent=styles['Normal'],
                textColor=text_dark, fontSize=9, fontName='Courier', leading=12,
                leftIndent=15, spaceAfter=8
            )
            for line in data["mindmap"].split("\n"):
                if line.strip():
                    story.append(Paragraph(line.replace(" ", "&nbsp;"), mindmap_style))
            story.append(Spacer(1, 4))

        # 9. Concept Glossary
        if data.get("concepts") and len(data["concepts"]) > 0:
            story.append(Paragraph("Concept Glossary", h2_style))
            for c_item in data["concepts"]:
                word = c_item.get("word", "")
                exp = c_item.get("explanation", "")
                story.append(Paragraph(f"<b>{word}</b>: {exp}", bullet_style))
            story.append(Spacer(1, 4))

        # 10. Sentence Importance Highlights
        if data.get("highlights") and len(data["highlights"]) > 0:
            story.append(Paragraph("Sentence Importance Highlights", h2_style))
            for h_item in data["highlights"]:
                h_text = h_item.get("text", "")
                importance = h_item.get("importance", "medium")
                label = "Very Important" if importance == "high" else "Moderately Important" if importance == "medium" else "Less Important"
                story.append(Paragraph(f"[{label}] {h_text}", bullet_style))

        # Footer
        story.append(Spacer(1, 30))
        story.append(HRFlowable(color=border_card, thickness=0.5, spaceAfter=8))
        story.append(Paragraph("Generated by SummrAI &bull; AI-Powered Email & Article Summarizer", footer_style))

        # Page numbering canvas callback
        def add_footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            canvas.setFillColor(colors.HexColor('#a0a0b0'))
            canvas.drawCentredString(letter[0]/2.0, 30, f"Page {doc.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
        return buffer.getvalue()

    except Exception as e:
        logger.warning(f"Styled PDF generation failed: {e}; falling back to TXT export")
        return export_as_txt(data)


def export_as_docx(data: dict) -> bytes:
    """Generate DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = data.get("title") or "SummrAI Summary"
        title_para = doc.add_heading(title, 0)
        title_para.runs[0].font.color.rgb = RGBColor(0x7C, 0x6A, 0xF7)

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        if data.get("summary"):
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(data["summary"])

        if data.get("bullets"):
            doc.add_heading("Key Points", level=1)
            for b in data["bullets"]:
                doc.add_paragraph(b, style='List Bullet')

        if data.get("takeaways"):
            doc.add_heading("Key Takeaways", level=1)
            for t in data["takeaways"]:
                doc.add_paragraph(t, style='List Bullet')

        if data.get("actions"):
            doc.add_heading("Action Items", level=1)
            for a in data["actions"]:
                doc.add_paragraph(a, style='List Number')

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        logger.warning("python-docx not installed; falling back to TXT export")
        return export_as_txt(data)
